import sys

if len(sys.argv) < 2:
    print("Usage: python your_script.py <file_path>")
    sys.exit(1)

file_path = sys.argv[1]

# Now you can use 'file_path' in your Python script
import aspose.words as aw
from docxtpl import DocxTemplate
from docx import Document
# import requests
from bs4 import BeautifulSoup
import chardet



# doc_tpl = DocxTemplate(output_file)

# context = {
#     'DRAFT REPORT' : ''
# }


# doc_tpl.render(context)
# print(doc_tpl.get_docx().toget_xml())
# doc_tpl.save(output_file)

# # # # print(tyoe(doc))




# # # To know the encoding 




# # from spire.doc import *

# from spire.doc.common import *
# doc_spire = Document()

# water = TextWatermark()
# water.text = ""

# doc_spire.LoadFromFile(output_file)









from docx.shared import Pt
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH




def update_document(doc_path, data_to_change):
    doc = Document(doc_path)

    for key, value in data_to_change.items():
        for paragraph in doc.paragraphs:
            if paragraph.text ==key:
                paragraph.text = value
                paragraph.text = paragraph.text.replace('\n', '')


    doc.save(doc_path)


# Parse HTML to find watermark and its CSS
def parse_html(html_file):
    with open(html_file, 'r', encoding='utf-8') as file:
        soup = BeautifulSoup(file, 'html.parser')
        watermark_element = soup.find('div', id='watermark') 
        # watermark_css = get_element_styles(watermark_element)
        watermark_css = ''

    return watermark_element, watermark_css

# # Extract styles from HTML element
# def get_element_styles(element):
#     if element:
#         style_string = element.get('style', '')
#         styles = dict(re.findall(r'(\S+?):(\S+?);', style_string))
#         return styles
#     else:
#         return {}

# Add watermark to Word document
def add_watermark_to_word(doc, watermark_text, styles):
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run(watermark_text)
        font = run.font
        font.size = Pt(int(40))



def add_page_numbers(doc):
    count =0
    total = 0
    for section in doc.sections:
        footer = section.footer

        # Check if there are paragraphs in the footer
        if footer:
            total +=1

    print(total)
    for section in doc.sections:
        count= count+1
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()



        run = paragraph.add_run()
        run.add_text("Page " + str(count) + " of " + str(total))
        run.add_text(str(count))
        # page_num_field = run.add_page_number()
        # run.add_text(" of ")
        # run.add_num_pages()


        # Adjust styles as needed
        font = run.font
        font.size = Pt(10)  # Adjust the font size
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align to the right


# if __name__ == "_main_":
    
input_file = file_path
# input_file = "uploads\sample_level_1.html"
input_file_name = input_file[8:-5]
output_file = "./WordDocs/"+input_file_name+".docx"

# import subprocess
# subprocess.run(['pandoc', input_file, "-o", output_file])


doc_aw = aw.Document(input_file)



doc_aw.save(output_file)



with open(input_file, 'r') as file:
    html_content = file.read()




soup_obj = BeautifulSoup(html_content, 'html.parser')
# print(soup_obj)
watermark_data = soup_obj.find('div', id = 'watermark')
# print(type(watermark_data.text))



# options = aw.TextWatermarkOptions()
# options.font_family = "Calibri"
# options.font_size = 42
# options.layout = aw.WatermarkLayout.DIAGONAL
# options.is_semitrasparent = True

# doc_aw.watermark.set_text(watermark_data.text, options)
# doc_aw.save(output_file)


doc_x  = Document(output_file)


# to remove the headers and footers
for section in doc_x.sections:
    header = section.header
    # section.different_first_page_header_footer = False
    # print(header)
    header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True




    # Add page numbers to the footer of each page
add_page_numbers(doc_x)

    # Save the modified Word document

doc_x.save(output_file)


        # Replace data in the document
data_to_change = {
        
    'Evaluation Only. Created with Aspose.Words. Copyright 2003-2023 Aspose Pty Ltd.' : ''
}

if watermark_data:
    data_to_change[watermark_data.text] = ''

update_document(output_file, data_to_change)

print("Document updated.")


    
watermark_element, watermark_css = parse_html(input_file)

if watermark_element:
        
    watermark_text = watermark_element.get_text()
    doc = Document(output_file)
    add_watermark_to_word(doc, watermark_text, watermark_css)
    doc.save(output_file)
else:
    print("Watermark not found in HTML.")











# input_file = file_path
# # input_file = "uploads\sample_level_1.html"
# input_file_name = input_file[8:-5]
# output_file = "./WordDocs/"+input_file_name+".docx"
# print(input_file)

# doc_aw = aw.Document(input_file)
# # print(type(doc_aw))

# doc_aw.save(output_file)

# with open(input_file, 'r') as file:
#     html_content = file.read()

# doc_x  = Document(output_file)

# # to remove the headers and footers
# for section in doc_x.sections:
#     header = section.header
#     # section.different_first_page_header_footer = False
#     print(header)
#     header.is_linked_to_previous = True
#     section.footer.is_linked_to_previous = True

# doc_x.save(output_file)

#print("Received file path:", file_path)


